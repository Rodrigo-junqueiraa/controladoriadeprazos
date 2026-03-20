"""Aba de Busca via DJEN"""

import tkinter as tk
from tkinter import messagebox
from datetime import datetime
from docx import Document
from ...core.configuracoes import TEMA
from ..styles import criar_botao_padrao, criar_label_padrao


class AbaSearchDJEN:
    """Aba para busca de publicações DJEN"""

    def __init__(self, parent, gerenciador_advogados):
        self.frame = tk.Frame(parent, bg=TEMA["bg_principal"])
        self.gerenciador_advogados = gerenciador_advogados
        self._criar_interface()
        self._atualizar_lista()

    def _criar_interface(self):
        """Cria elementos da interface"""
        # Frame principal
        frame_adv = tk.Frame(self.frame, bg=TEMA["bg_principal"])
        frame_adv.pack(pady=10)

        # Entrada de nome
        label_nome = criar_label_padrao(frame_adv, "Nome do Advogado:")
        label_nome.pack()

        self.entrada_nome_adv = tk.Entry(frame_adv, width=50)
        self.entrada_nome_adv.pack(pady=5)

        # Listbox de advogados
        self.lista_advogados = tk.Listbox(frame_adv, width=50, height=8)
        self.lista_advogados.pack(pady=5)

        # Botões
        btn_add = criar_botao_padrao(
            frame_adv,
            "Adicionar Advogado",
            self._adicionar_advogado
        )
        btn_add.pack(pady=3)

        btn_del = criar_botao_padrao(
            frame_adv,
            "Excluir Advogado Selecionado",
            self._excluir_advogado,
            cor_bg=TEMA["cor_vermelho"]
        )
        btn_del.pack(pady=3)

        btn_buscar = criar_botao_padrao(
            frame_adv,
            "Buscar Publicações do Dia",
            self._buscar_publicacoes,
            cor_bg=TEMA["cor_azul"]
        )
        btn_buscar.pack(pady=10)

    def _atualizar_lista(self):
        """Atualiza listbox com advogados memorados"""
        self.lista_advogados.delete(0, tk.END)
        for nome in self.gerenciador_advogados.carregar():
            self.lista_advogados.insert(tk.END, nome)

    def _adicionar_advogado(self):
        """Adiciona novo advogado"""
        nome = self.entrada_nome_adv.get().strip()

        if not nome:
            messagebox.showwarning("Aviso", "Digite o nome do advogado.")
            return

        self.lista_advogados.insert(tk.END, nome)

        resposta = messagebox.askyesno(
            "Memorizar?",
            f"Deseja que o sistema memorize o nome '{nome}' para futuras buscas?"
        )

        if resposta:
            self.gerenciador_advogados.adicionar(nome)
            self._atualizar_lista()

        self.entrada_nome_adv.delete(0, tk.END)

    def _excluir_advogado(self):
        """Excluir advogado da lista memorada"""
        selecionado = self.lista_advogados.curselection()

        if not selecionado:
            messagebox.showwarning("Aviso", "Selecione um nome para excluir.")
            return

        index = selecionado[0]
        nome = self.lista_advogados.get(index)
        self.lista_advogados.delete(index)
        self.gerenciador_advogados.remover(nome)

    def _buscar_publicacoes(self):
        """Busca publicações (simulado)"""
        advogados = self.gerenciador_advogados.carregar()

        if not advogados:
            messagebox.showwarning("Aviso", "Nenhum advogado cadastrado para busca.")
            return

        # Dados simulados
        publicacoes = [
            {
                "data_publicacao": "29/04/2025",
                "data_disponibilizacao": "28/04/2025",
                "processo": "0010393-96.2020.5.03.0026",
                "variacao": "RODRIGO JUNQUEIRA DE LIMA SIQUEIRA",
                "tribunal": "TRT3 - MINAS GERAIS",
                "orgao": "1ª Vara do Trabalho de Betim",
                "conteudo": "Publicação simulada com prazo.",
                "link": "https://exemplo.com/publicacao1"
            },
            {
                "data_publicacao": "29/04/2025",
                "data_disponibilizacao": "28/04/2025",
                "processo": "0001234-56.2024.5.10.0001",
                "variacao": "DENISE MARCONDES",
                "tribunal": "TRT10 - DISTRITO FEDERAL",
                "orgao": "2ª Vara do Trabalho de Brasília",
                "conteudo": "Publicação do TRT10 para manifestação.",
                "link": "https://exemplo.com/publicacao2"
            }
        ]

        # Separar por tipo
        mg = [p for p in publicacoes if "MG" in p["tribunal"] or "Minas Gerais" in p["tribunal"]]
        outros = [p for p in publicacoes if "MG" not in p["tribunal"] and "Minas Gerais" not in p["tribunal"]]

        # Gerar documentos
        self._gerar_docx(mg, f"JORNAL_MG_{datetime.now().strftime('%d-%m-%Y - %Hh%M')}.docx")
        self._gerar_docx(outros, f"JORNAL_NACIONAL_{datetime.now().strftime('%d-%m-%Y - %Hh%M')}.docx")

        messagebox.showinfo("Sucesso", "Arquivos DOCX gerados com sucesso!")

    def _gerar_docx(self, publicacoes, nome_arquivo):
        """Gera arquivo DOCX com publicações"""
        doc = Document()

        for pub in publicacoes:
            doc.add_paragraph(f"DATA DA PUBLICAÇÃO: {pub['data_publicacao']}    DATA DISP: {pub['data_disponibilizacao']}")
            doc.add_paragraph(f"PROCESSO: {pub['processo']}")
            doc.add_paragraph(f"VARIAÇÃO: {pub['variacao']}")
            doc.add_paragraph(f"TRIBUNAL: {pub['tribunal']}")
            doc.add_paragraph(f"ÓRGÃO: {pub['orgao']}")
            doc.add_paragraph(f"Conteúdo:\n{pub['conteudo']}")
            doc.add_paragraph(f"Link:\n{pub['link']}")
            doc.add_paragraph("-" * 60)

        doc.save(nome_arquivo)

    def get_frame(self):
        """Retorna o frame da aba"""
        return self.frame
